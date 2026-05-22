"""개발 제안서 Markdown → DOCX보내기."""

from __future__ import annotations

import io
import re
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^[-*+]\s+(.*)$")
_NUM = re.compile(r"^\d+\.\s+(.*)$")
_INLINE_MD = re.compile(r"\*\*(.+?)\*\*|`([^`]+)`|\*(.+?)\*")


def _plain_line(line: str) -> str:
    text = line.strip()

    def _sub(m: re.Match[str]) -> str:
        return m.group(1) or m.group(2) or m.group(3) or ""

    while True:
        new = _INLINE_MD.sub(_sub, text)
        if new == text:
            break
        text = new
    return text


def proposal_markdown_to_docx_bytes(markdown_text: str, *, document_title: str = "") -> bytes:
    from docx import Document

    doc = Document()
    title = (document_title or "").strip()
    if title:
        doc.add_heading(title, level=0)

    for raw in (markdown_text or "").splitlines():
        stripped = raw.strip()
        if not stripped:
            continue
        hm = _HEADING.match(stripped)
        if hm:
            level = min(len(hm.group(1)), 4)
            doc.add_heading(_plain_line(hm.group(2)), level=level)
            continue
        bm = _BULLET.match(stripped)
        if bm:
            doc.add_paragraph(_plain_line(bm.group(1)), style="List Bullet")
            continue
        nm = _NUM.match(stripped)
        if nm:
            doc.add_paragraph(_plain_line(nm.group(1)), style="List Number")
            continue
        doc.add_paragraph(_plain_line(stripped))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def proposal_download_filename(
    request_kind: str, request_id: int, *, fmt: str, title: str | None = None
) -> str:
    base = re.sub(r"[^\w\-]+", "_", (title or "").strip())[:40].strip("_") or "proposal"
    ext = "docx" if fmt == "docx" else "md"
    kind = (request_kind or "rfp").strip().lower()
    return f"{base}_{kind}_{request_id}.{ext}"
